import io
import re
import uuid
from typing import Any, Dict, List, Tuple
from langchain.tools import tool
from docx import Document
from docx.shared import Pt

from services.context import get_current_user_id
from services.files import FileStore
from services.ratelimit import get_rate_limiter
from services.storage import StorageError


@tool
def create_docx(title: str, content: str) -> str:
    """Create a Word document (.docx file).

    Use ONLY when the user explicitly asks for a document, essay, report,
    resume, or Word file. Do NOT use for chat responses, quick answers,
    or when the user just wants information.

    Args:
        title: Document heading.
        content: Body text; each newline-separated line becomes a paragraph.

    Returns:
        Filename of the saved document (plus its download ID).
    """
    try:
        doc: Document = Document()
        doc.add_heading(title, level=1)
        for paragraph in content.strip().split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        filename: str = f"docx_{uuid.uuid4().hex[:8]}.docx"
        buf = io.BytesIO()
        doc.save(buf)
        data: bytes = buf.getvalue()

        user_id = get_current_user_id()
        if user_id:
            verdict = get_rate_limiter().check(user_id, "generate")
            if not verdict.allowed:
                return (
                    "STATUS=DENIED tool=create_docx: generation rate limit "
                    f"exceeded, retry in {verdict.retry_after:.0f}s."
                )
            try:
                meta = FileStore(user_id).register_output(filename, data, "docx")
                return f"Document saved as {meta.display_name} (file ID: {meta.id})"
            except StorageError as e:
                return f"STATUS=FAILED tool=create_docx: {e}"
        with open(filename, "wb") as f:
            f.write(data)
        return f"Document saved as {filename}"
    except Exception as e:
        return f"STATUS=FAILED tool=create_docx: {str(e)}"


_DOCX_MAX_BLOCKS: int = 300
_DOCX_MAX_CELL_CHARS: int = 200
_DOCX_MAX_TABLE_COLS: int = 8


def _docx_add_runs(paragraph: Any, text: str) -> None:
    """Add **bold**, *italic*, and `code` inline runs to a paragraph."""
    token_re = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    pos = 0
    first = True
    for match in token_re.finditer(text):
        if match.start() > pos:
            _docx_plain_run(paragraph, text[pos:match.start()], first)
            first = False
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2]) if not first else _docx_first_run(paragraph, token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1]) if not first else _docx_first_run(paragraph, token[1:-1])
            try:
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            except Exception:
                pass
        else:
            run = paragraph.add_run(token[1:-1]) if not first else _docx_first_run(paragraph, token[1:-1])
            run.italic = True
        first = False
        pos = match.end()
    if pos < len(text):
        _docx_plain_run(paragraph, text[pos:], first)


def _docx_first_run(paragraph: Any, text: str) -> Any:
    """Reuse the paragraph's initial empty run when possible."""
    try:
        if paragraph.runs:
            paragraph.runs[0].text = text
            return paragraph.runs[0]
    except Exception:
        pass
    return paragraph.add_run(text)


def _docx_plain_run(paragraph: Any, text: str, first: bool) -> None:
    """Append plain text, reusing the first run when available."""
    if not text:
        return
    if first:
        _docx_first_run(paragraph, text)
    else:
        paragraph.add_run(text)


def _docx_parse_blocks(markdown_text: str) -> List[Tuple[str, Any]]:
    """Parse lightweight markdown into (kind, payload) blocks.

    Kinds: h1/h2/h3, para, bullet, numbered, quote, code, table, pagebreak.
    """
    blocks: List[Tuple[str, Any]] = []
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    in_code = False
    code_lang = ""
    code_buf: List[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                blocks.append(("code", {"lang": code_lang, "code": "\n".join(code_buf)}))
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_lang = stripped[3:].strip()[:20]
            i += 1
            continue
        if in_code:
            code_buf.append(line.rstrip())
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped in ("---", "***", "___"):
            blocks.append(("pagebreak", None))
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append(("h3", stripped[4:].strip()))
        elif stripped.startswith("## "):
            blocks.append(("h2", stripped[3:].strip()))
        elif stripped.startswith("# "):
            blocks.append(("h1", stripped[2:].strip()))
        elif stripped.startswith("> "):
            quote_lines = [stripped[2:].strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(("quote", " ".join(quote_lines)))
            continue
        elif re.match(r"^[-*]\s+", stripped):
            items = [re.sub(r"^[-*]\s+", "", stripped).strip()]
            i += 1
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()).strip())
                i += 1
            blocks.append(("bullet", [x for x in items if x]))
            continue
        elif re.match(r"^\d+[.)]\s+", stripped):
            items = [re.sub(r"^\d+[.)]\s+", "", stripped).strip()]
            i += 1
            while i < len(lines) and re.match(r"^\d+[.)]\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+[.)]\s+", "", lines[i].strip()).strip())
                i += 1
            blocks.append(("numbered", [x for x in items if x]))
            continue
        elif "|" in stripped:
            rows = [stripped]
            i += 1
            while i < len(lines) and "|" in lines[i]:
                rows.append(lines[i].strip())
                i += 1
            parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
            parsed = [r for r in parsed if any(c for c in r)]
            # Drop the markdown separator row (| --- | --- |).
            parsed = [r for r in parsed if not all(re.fullmatch(r":?-{1,}:?", c or "") for c in r)]
            if len(parsed) >= 1 and len(parsed[0]) >= 1 and len(parsed[0]) <= _DOCX_MAX_TABLE_COLS:
                blocks.append(("table", parsed))
            else:
                for r in parsed:
                    blocks.append(("para", " | ".join(r)))
            continue
        else:
            para_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip() and not re.match(
                r"^(#{1,3}\s|>\s|[-*]\s|\d+[.)]\s|```|---$|\*\*\*$|___$)", lines[i].strip()
            ) and "|" not in lines[i]:
                para_lines.append(lines[i].strip())
                i += 1
            blocks.append(("para", " ".join(para_lines)))
            continue
        i += 1
    if in_code and code_buf:
        blocks.append(("code", {"lang": code_lang, "code": "\n".join(code_buf)}))
    return blocks[:_DOCX_MAX_BLOCKS]


@tool
def build_document(markdown_text: str, title: str = "Document") -> str:
    """Build a structured Word document from lightweight markdown.

    Prefer this over create_docx when structure matters. Supported:
    # / ## / ### headings, paragraphs, - bullets, 1. numbered lists,
    > quotes, ``` code blocks, | tables |, --- page breaks, and
    **bold** / *italic* / `code` inline runs.

    Args:
        markdown_text: The document body in lightweight markdown.
        title: Document title (becomes Heading 1 + filename basis).

    Returns:
        Summary with filename (plus download ID) and section counts,
        or a STATUS= error marker.
    """
    if not markdown_text or not markdown_text.strip():
        return "STATUS=INVALID tool=build_document: empty document text."
    if not title or not title.strip():
        return "STATUS=INVALID tool=build_document: empty title."
    try:
        blocks = _docx_parse_blocks(markdown_text)
        if not any(kind in ("para", "bullet", "numbered", "table", "code", "quote", "h2", "h3") for kind, _ in blocks):
            return "STATUS=INVALID tool=build_document: no substantive content found."
        doc: Document = Document()
        try:
            style = doc.styles["Normal"]
            style.font.size = Pt(11)
        except Exception:
            pass
        doc.add_heading(title.strip()[:120], level=1)
        counts: Dict[str, int] = {}
        for kind, payload in blocks:
            counts[kind] = counts.get(kind, 0) + 1
            if kind in ("h1", "h2", "h3"):
                doc.add_heading(str(payload)[:200], level={"h1": 1, "h2": 2, "h3": 3}[kind])
            elif kind == "para":
                para = doc.add_paragraph()
                _docx_add_runs(para, str(payload)[:3000])
            elif kind == "bullet":
                for item in payload[:50]:
                    para = doc.add_paragraph(style="List Bullet")
                    _docx_add_runs(para, str(item)[:1000])
            elif kind == "numbered":
                for item in payload[:50]:
                    para = doc.add_paragraph(style="List Number")
                    _docx_add_runs(para, str(item)[:1000])
            elif kind == "quote":
                try:
                    para = doc.add_paragraph(style="Intense Quote")
                except Exception:
                    para = doc.add_paragraph()
                _docx_add_runs(para, str(payload)[:2000])
            elif kind == "code":
                para = doc.add_paragraph()
                try:
                    para.style = doc.styles["No Spacing"]
                except Exception:
                    pass
                run = para.add_run(str(payload.get("code", ""))[:4000])
                try:
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                except Exception:
                    pass
            elif kind == "table":
                headers = payload[0]
                rows = payload[1:13]
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:
                    pass
                for j, header in enumerate(headers):
                    table.cell(0, j).text = str(header)[:_DOCX_MAX_CELL_CHARS]
                for i, row in enumerate(rows):
                    for j in range(len(headers)):
                        table.cell(i + 1, j).text = str(row[j] if j < len(row) else "")[:_DOCX_MAX_CELL_CHARS]
            elif kind == "pagebreak":
                try:
                    doc.add_page_break()
                except Exception:
                    pass
        filename: str = f"docx_{uuid.uuid4().hex[:8]}.docx"
        buf = io.BytesIO()
        doc.save(buf)
        data: bytes = buf.getvalue()

        # Quality control: reopen and verify structure survived.
        try:
            check = Document(io.BytesIO(data))
            paras = [p for p in check.paragraphs if (p.text or "").strip()]
            tables = list(check.tables)
            expected_tables = sum(1 for k, _ in blocks if k == "table")
            if not paras:
                return "STATUS=FAILED tool=build_document: validation failed (no readable paragraphs)."
            if len(tables) != expected_tables:
                return (
                    "STATUS=FAILED tool=build_document: validation failed "
                    f"(expected {expected_tables} tables, found {len(tables)})."
                )
        except Exception as e:
            return f"STATUS=FAILED tool=build_document: output unreadable ({str(e)[:120]})."

        summary = (
            f"Document saved as {filename} with "
            f"{counts.get('h1', 0) + counts.get('h2', 0) + counts.get('h3', 0)} headings, "
            f"{counts.get('para', 0)} paragraphs, "
            f"{counts.get('bullet', 0) + counts.get('numbered', 0)} lists, "
            f"{counts.get('table', 0)} tables."
        )
        user_id = get_current_user_id()
        if user_id:
            verdict = get_rate_limiter().check(user_id, "generate")
            if not verdict.allowed:
                return (
                    "STATUS=DENIED tool=build_document: generation rate limit "
                    f"exceeded, retry in {verdict.retry_after:.0f}s."
                )
            try:
                meta = FileStore(user_id).register_output(filename, data, "docx")
                return f"{summary} (file ID: {meta.id})"
            except StorageError as e:
                return f"STATUS=FAILED tool=build_document: {e}"
        with open(filename, "wb") as f:
            f.write(data)
        return summary
    except Exception as e:
        return f"STATUS=FAILED tool=build_document: {str(e)[:200]}"
